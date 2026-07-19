"""
Self-Study Report (SSR) generator — final integration across all 7 standards.

Builds one Word document, in Arabic with right-to-left layout (matching the
official NAQAAE-style indicator wording already used elsewhere in the app),
with a section per standard, pulling automatically from each phase's own
storage:

  Standard 1  governance.py            mission text, document register, stakeholder log
  Standard 2  curriculum_mapping.py    ILOs, coverage-matrix findings
  Standard 3  academic_analytics/kpis  KPI figures, via the same derivation program_report_docx.py
                                       already uses (reused here, not reimplemented)
  Standard 4  alumni.py                graduate roster stats (+ a pointer to the survey dashboard,
                                       since survey responses are anonymous/aggregate — see alumni.py)
  Standard 5  faculty_data.py          roster, load-imbalance / specialization-gap flags
  Standard 6  resources.py             equipment/library/budget, maintenance-due
  Standard 7  indicators.py            per-standard status + closing-the-loop log

Built from scratch with python-docx, following the same helper pattern as
course_report_docx.py and curriculum_map_report.py (not the XML-template
approach in program_report_docx.py, since there's no fixed NAQAAE SSR
template to fill in here — only its KPI-derivation helper is reused).

RTL note: Word renders Arabic text and shapes it correctly on its own —
unlike the matplotlib charts elsewhere in the app, no arabic_reshaper/
python-bidi pre-shaping is needed here. What Word needs instead is explicit
direction markup (w:bidi on paragraphs, w:bidiVisual on tables) plus a
complex-script font (w:rFonts/w:cs), set once on the base styles so every
paragraph/table built through the helpers below inherits it automatically.
"""
from __future__ import annotations

from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

import alumni
import curriculum_mapping
import faculty_data
import governance
import indicators
import resources

_HEADER_BLUE = RGBColor(0x1E, 0x3A, 0x5F)
_LIGHT_GREY = "D9D9D9"
_ARABIC_FONT = "Arial"
_ARABIC_LANG = "ar-EG"

_STANDARD_NAMES_AR = {
    1: "رسالة وإدارة البرنامج",
    2: "تصميم البرنامج",
    3: "التعليم والتعلم والتقييم",
    4: "الطلاب والخريجون",
    5: "أعضاء هيئة التدريس والهيئة المعاونة",
    6: "الموارد ومصادر التعلم والتسهيلات الداعمة",
    7: "ضمان الجودة وتقييم البرنامج",
}


# ---------------------------------------------------------------------------
# RTL / Arabic docx helpers
# ---------------------------------------------------------------------------

def _make_rtl_paragraph_props(pPr) -> None:
    if pPr.find(qn("w:bidi")) is None:
        pPr.append(OxmlElement("w:bidi"))


def _make_rtl_run_props(rPr) -> None:
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:cs"), _ARABIC_FONT)
    lang = rPr.find(qn("w:lang"))
    if lang is None:
        lang = OxmlElement("w:lang")
        rPr.append(lang)
    lang.set(qn("w:bidi"), _ARABIC_LANG)


def _apply_rtl_style(style) -> None:
    """Makes every paragraph/run built with this style RTL by default, so
    individual add_paragraph() calls throughout the section builders don't
    each need their own direction markup."""
    style.font.name = _ARABIC_FONT
    style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _make_rtl_paragraph_props(style.element.get_or_add_pPr())
    _make_rtl_run_props(style.element.get_or_add_rPr())


def _set_table_rtl(table) -> None:
    tblPr = table._tbl.tblPr
    if tblPr.find(qn("w:bidiVisual")) is None:
        tblPr.append(OxmlElement("w:bidiVisual"))
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT


# ---------------------------------------------------------------------------
# Shared docx helpers (same pattern as course_report_docx.py / curriculum_map_report.py)
# ---------------------------------------------------------------------------

def _shade_cell(cell, hex_color: str) -> None:
    shd = cell._tc.get_or_add_tcPr().makeelement(qn("w:shd"), {
        qn("w:val"): "clear", qn("w:color"): "auto", qn("w:fill"): hex_color,
    })
    cell._tc.get_or_add_tcPr().append(shd)


def _bold(cell, text: str, size: int | None = None) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.bold = True
    if size is not None:
        run.font.size = Pt(size)


def _standard_heading(doc: Document, number: int, title: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(20)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(f"المعيار {number} — {title}")
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = _HEADER_BLUE


def _subheading(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(11.5)


def _bullet_list(doc: Document, items: list[str], empty_text: str) -> None:
    if not items:
        doc.add_paragraph().add_run(empty_text).italic = True
        return
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def _simple_table(doc: Document, headers: list[str], rows: list[list[str]], empty_text: str) -> None:
    if not rows:
        doc.add_paragraph().add_run(empty_text).italic = True
        return
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        _bold(hdr[i], h, size=9.5)
        _shade_cell(hdr[i], _LIGHT_GREY)
    for row_vals in rows:
        cells = table.add_row().cells
        for i, v in enumerate(row_vals):
            cells[i].text = str(v)
    _set_table_rtl(table)


# ---------------------------------------------------------------------------
# Section builders — one per standard
# ---------------------------------------------------------------------------

def _build_standard_1(doc: Document) -> None:
    _standard_heading(doc, 1, _STANDARD_NAMES_AR[1])

    _subheading(doc, "رسالة البرنامج")
    mission = governance.get_current_mission()
    doc.add_paragraph(mission["mission_text"] if mission else "لم يتم تسجيل نص الرسالة بعد.")
    versions = governance.list_mission_versions()
    if len(versions) > 1:
        doc.add_paragraph(f"(يوجد {len(versions)} إصدارات مسجلة؛ يظهر أعلاه أحدثها.)").italic = True

    _subheading(doc, "سجل مستندات المجالس واللجان")
    docs = governance.list_documents()
    _simple_table(
        doc,
        ["العنوان", "اللجنة", "التاريخ"],
        [[d["title"], d.get("committee_name") or "", d.get("document_date") or ""] for d in docs],
        "لم يتم تسجيل أي مستندات حوكمة بعد.",
    )

    _subheading(doc, "سجل مشاركة أصحاب المصلحة")
    log = governance.list_stakeholder_log()
    _simple_table(
        doc,
        ["صاحب المصلحة", "الدور", "التاريخ", "الموضوع"],
        [[e["stakeholder_name"], e.get("stakeholder_role") or "", e["consulted_on"], e["topic"]] for e in log],
        "لم يتم تسجيل أي استشارات لأصحاب المصلحة بعد.",
    )


def _build_standard_2(doc: Document) -> None:
    _standard_heading(doc, 2, _STANDARD_NAMES_AR[2])

    ilos = curriculum_mapping.list_ilos()
    courses = curriculum_mapping.list_courses()
    summary = curriculum_mapping.compute_coverage_summary()

    doc.add_paragraph(
        f"تم ربط {len(ilos)} من مخرجات التعلم المستهدفة للبرنامج (ILOs) عبر {len(courses)} مقررًا. "
        f"مصفوفة التغطية الكاملة (المقررات × المخرجات) متاحة عبر تصدير DOCX من صفحة خريطة المنهج."
    )

    _subheading(doc, "مخرجات التعلم المستهدفة للبرنامج")
    _simple_table(
        doc,
        ["الرمز", "النص"],
        [[i.get("ilo_code") or f"ILO{i['id']}", i["ilo_text"]] for i in ilos],
        "لم يتم إدخال أي مخرجات تعلم بعد.",
    )

    _subheading(doc, "نتائج التغطية")
    doc.add_paragraph("مخرجات التعلم غير المغطاة:")
    _bullet_list(
        doc,
        [f"{i.get('ilo_code') or ('ILO' + str(i['id']))}: {i['ilo_text']}" for i in summary["zero_coverage_ilos"]],
        "لا يوجد — كل مخرج تعلم مغطى بمقرر واحد على الأقل.",
    )
    doc.add_paragraph("المقررات غير المرتبطة بأي مخرج تعلم:")
    _bullet_list(
        doc,
        [c["course_name"] for c in summary["courses_without_ilos"]],
        "لا يوجد — كل مقرر مرتبط بمخرج تعلم واحد على الأقل.",
    )


def _build_standard_3(doc: Document, analysis: dict[str, Any] | None) -> None:
    _standard_heading(doc, 3, _STANDARD_NAMES_AR[3])

    if not analysis:
        doc.add_paragraph(
            "لم يتم توفير بيانات تحليلية لهذا التقرير. لإنشاء هذا القسم، قم برفع ملف الدرجات في "
            "لوحة الجودة ثم أعد تصدير تقرير الدراسة الذاتية مع إرفاق نتائج التحليل."
        ).italic = True
        return

    from program_report_docx import _derive_program_metrics, _fmt_gpa, _fmt_number, _fmt_percent

    metrics = _derive_program_metrics(analysis)
    _simple_table(
        doc,
        ["المؤشر", "القيمة"],
        [
            ["عدد المقررات التي تم تحليلها", _fmt_number(metrics["course_count"])],
            ["إجمالي عدد الطلاب المسجلين", _fmt_number(metrics["total_enrollment"])],
            ["متوسط المعدل التراكمي", _fmt_gpa(metrics["avg_gpa"])],
            ["متوسط نسبة الرسوب", _fmt_percent(metrics["failure_rate"])],
            ["متوسط نسبة التفوق", _fmt_percent(metrics["excellence_rate"])],
            ["عدد المقررات التي تتجاوز فيها نسبة الرسوب 20%", _fmt_number(metrics["courses_over_20_failure"])],
        ],
        "لا توجد بيانات مؤشرات أداء متاحة.",
    )


def _build_standard_4(doc: Document) -> None:
    _standard_heading(doc, 4, _STANDARD_NAMES_AR[4])

    summary = alumni.get_registry_summary()
    _simple_table(
        doc,
        ["المؤشر", "القيمة"],
        [
            ["عدد الخريجين المسجلين", str(summary["total_alumni"])],
            ["نسبة التوظيف", f"{summary['employment_rate']}%"],
            ["نسبة المشاركة في الاستبيان", f"{summary['survey_participation_rate']}%"],
        ],
        "لا توجد بيانات خريجين متاحة.",
    )
    doc.add_paragraph(
        "ملاحظة: استجابات الاستبيان التي تم جمعها عبر لوحة الاستبيان مجهولة الهوية وتجميعية فقط، لذا لا "
        "يمكن ربطها بسجلات الخريجين الفردية هنا. يُرجى إرفاق نتائج الرضا من لوحة الاستبيان (تصدير "
        "PPTX/PNG) بشكل منفصل كدليل داعم لهذا المعيار."
    ).italic = True


def _build_standard_5(doc: Document) -> None:
    _standard_heading(doc, 5, _STANDARD_NAMES_AR[5])

    summary = faculty_data.get_dashboard_summary()
    _simple_table(
        doc,
        ["المؤشر", "القيمة"],
        [
            ["عدد أعضاء هيئة التدريس", str(summary["total_faculty"])],
            ["متوسط العبء التدريسي (ساعة/فصل دراسي)", str(summary["average_load_hours"])],
            ["حالات زيادة العبء", str(summary["overloaded_count"])],
            ["حالات نقص العبء", str(summary["underloaded_count"])],
            ["فجوات التخصص", str(summary["specialization_gap_count"])],
        ],
        "لا توجد بيانات لأعضاء هيئة التدريس.",
    )

    _subheading(doc, "فجوات التخصص")
    _bullet_list(
        doc,
        [f"{g['course_name']} ({g['semester']}) — {g['faculty_name']}" for g in summary["specialization_gap_flags"]],
        "لم يتم رصد أي فجوات.",
    )


def _build_standard_6(doc: Document) -> None:
    _standard_heading(doc, 6, _STANDARD_NAMES_AR[6])

    summary = resources.get_dashboard_summary()
    _simple_table(
        doc,
        ["المؤشر", "القيمة"],
        [
            ["عدد الأجهزة والمعدات", str(summary["total_equipment"])],
            ["بحاجة إلى صيانة", str(summary["needs_repair_count"])],
            ["خارج الخدمة", str(summary["out_of_service_count"])],
            ["صيانة مستحقة (خلال 30 يومًا)", str(summary["maintenance_due_count"])],
            ["مقتنيات المكتبة", f"{summary['total_library_items']} ({summary['total_library_titles']} عنوان)"],
            ["إجمالي الميزانية المسجلة", str(summary["total_budget_amount"])],
        ],
        "لا توجد بيانات موارد متاحة.",
    )

    _subheading(doc, "الصيانة المستحقة")
    _simple_table(
        doc,
        ["الجهاز", "الموقع", "تاريخ الاستحقاق", "الحالة"],
        [
            [e["name"], e.get("location") or "", e["next_maintenance_date"], "متأخر" if e["overdue"] else "مستحق قريبًا"]
            for e in summary["maintenance_due"]
        ],
        "لا توجد صيانة مستحقة خلال الثلاثين يومًا القادمة.",
    )


def _build_standard_7(doc: Document) -> None:
    _standard_heading(doc, 7, _STANDARD_NAMES_AR[7])

    _subheading(doc, "حالة المؤشرات حسب المعيار")
    standard_summaries = indicators.summarize_by_standard()
    _simple_table(
        doc,
        ["المعيار", "الإجمالي", "مكتمل", "جزئي", "غير موجود"],
        [
            [f"{s.standard_number}. {_STANDARD_NAMES_AR.get(s.standard_number, s.standard_name)}",
             str(s.total), str(s.complete), str(s.partial), str(s.missing)]
            for s in standard_summaries
        ],
        "لم يتم العثور على أي مؤشرات.",
    )

    _subheading(doc, "إغلاق حلقة التحسين")
    all_indicators = indicators.list_indicators()
    loop_rows: list[list[str]] = []
    for ind in all_indicators:
        full = indicators.get_indicator(ind["id"])
        for entry in full.get("closing_the_loop_log", []):
            loop_rows.append([
                f"المعيار {ind['standard_number']}: {ind['indicator_text'][:60]}",
                entry["weakness_identified"],
                entry.get("action_taken") or "",
                entry["entry_date"],
                entry.get("entry_status") or "",
            ])
    _simple_table(
        doc,
        ["المؤشر", "الضعف الذي تم رصده", "الإجراء المتخذ", "التاريخ", "الحالة"],
        loop_rows,
        "لم يتم تسجيل أي إدخالات لإغلاق حلقة التحسين بعد.",
    )


# ---------------------------------------------------------------------------
# Public entry point
# ---------------------------------------------------------------------------

def build_ssr_docx(analysis: dict[str, Any] | None = None) -> bytes:
    import io

    doc = Document()
    for style_name in ("Normal", "List Bullet"):
        _apply_rtl_style(doc.styles[style_name])
    doc.styles["Normal"].font.size = Pt(10.5)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("تقرير الدراسة الذاتية")
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = _HEADER_BLUE

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.add_run(
        "تم إعداده تلقائيًا من وحدات دعم الاعتماد الخاصة بالبرنامج "
        "(الحوكمة، خريطة المنهج، أعضاء هيئة التدريس، الموارد، الخريجون، ومتابع المؤشرات)."
    ).italic = True

    _build_standard_1(doc)
    _build_standard_2(doc)
    _build_standard_3(doc, analysis)
    _build_standard_4(doc)
    _build_standard_5(doc)
    _build_standard_6(doc)
    _build_standard_7(doc)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
