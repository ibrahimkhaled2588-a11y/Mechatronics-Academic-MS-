"""
Course Report DOCX builder.

Produces a Word document that follows the standard NAQAAE-style "Course Report"
template (Basic Information / Data and Statistics / Student Assessment Results /
Course Quality Evaluation / Student Feedback / Instructors' Reflection /
Course Enhancement / Signatures), built from scratch with python-docx and filled
with the indicators already computed by the analytics pipeline for the selected
course (enrollment, grade distribution, pass/fail counts, GPA, recommendations),
plus any manually supplied basic-info fields (course code, instructor, dates, ...)
that cannot be derived from an uploaded grades workbook.
"""
from __future__ import annotations

from typing import Any

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

_HEADER_BLUE = RGBColor(0x1E, 0x3A, 0x5F)
_LIGHT_GREY = "D9D9D9"

GRADE_ORDER = ["A+", "A", "A-", "B+", "B", "B-", "C+", "C", "C-", "D+", "D", "F"]


# ---------------------------------------------------------------------------
# Small helpers
# ---------------------------------------------------------------------------

def _safe_float(value: Any) -> float | None:
    try:
        n = float(value)
    except (TypeError, ValueError):
        return None
    return None if n != n else n


def _fmt_num(value: Any) -> str:
    n = _safe_float(value)
    return "Not available" if n is None else f"{int(round(n)):,}"


def _fmt_pct(value: Any, digits: int = 1) -> str:
    n = _safe_float(value)
    return "Not available" if n is None else f"{n:.{digits}f}%"


def _fmt_gpa(value: Any) -> str:
    n = _safe_float(value)
    return "Not available" if n is None else f"{n:.2f} / 4.00"


def _shade_cell(cell, hex_color: str) -> None:
    shd = cell._tc.get_or_add_tcPr().makeelement(qn("w:shd"), {
        qn("w:val"): "clear", qn("w:color"): "auto", qn("w:fill"): hex_color,
    })
    cell._tc.get_or_add_tcPr().append(shd)


def _bold(cell, text: str, color: RGBColor | None = None, size: int | None = None) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.bold = True
    if color is not None:
        run.font.color.rgb = color
    if size is not None:
        run.font.size = Pt(size)


def _section_heading(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(13)
    run.font.color.rgb = _HEADER_BLUE


def _note(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)


def _label_value_table(doc: Document, rows: list[tuple[str, str]]) -> None:
    table = doc.add_table(rows=0, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for label, value in rows:
        row = table.add_row()
        _bold(row.cells[0], label)
        row.cells[1].text = value or "Not available"


# ---------------------------------------------------------------------------
# Data shaping
# ---------------------------------------------------------------------------

def _grade_counts_from_pct(grade_distribution: dict[str, Any], enrollment: float | None) -> dict[str, int]:
    """Approximate grade counts from percentages + total enrollment."""
    counts: dict[str, int] = {}
    total = _safe_float(enrollment) or 0.0
    for g in GRADE_ORDER:
        pct = _safe_float((grade_distribution or {}).get(g))
        counts[g] = int(round((pct or 0.0) / 100.0 * total)) if total > 0 else 0
    return counts


def _derive_assessment_results(payload: dict[str, Any]) -> dict[str, Any]:
    enrollment = _safe_float(payload.get("enrollment")) or 0.0
    failure_rate = _safe_float(payload.get("failure_rate"))
    grade_dist = payload.get("grade_distribution") or {}
    grade_counts = _grade_counts_from_pct(grade_dist, enrollment)
    fail_count = grade_counts.get("F", 0)
    # Prefer the failure_rate KPI (Wilson-adjusted, matches dashboard) if grade-letter F count looks off
    if failure_rate is not None and enrollment > 0:
        fail_count_from_rate = int(round(failure_rate / 100.0 * enrollment))
        fail_count = fail_count_from_rate
    passed = max(0, int(round(enrollment)) - fail_count)
    success_pct = (passed / enrollment * 100.0) if enrollment > 0 else None
    return {
        "enrollment": enrollment,
        "grade_counts": grade_counts,
        "fail_count": fail_count,
        "passed": passed,
        "success_pct": success_pct,
        "failure_pct": failure_rate,
    }


# ---------------------------------------------------------------------------
# Section builders
# ---------------------------------------------------------------------------

def _build_title(doc: Document, payload: dict[str, Any]) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"Course Report ({payload.get('academic_year') or '2025'})")
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = _HEADER_BLUE

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = sub.add_run(
        f"Academic Year: {payload.get('academic_year') or 'Not available'}   |   "
        f"Semester: {payload.get('semester') or 'Not available'}"
    )
    sub_run.bold = True


def _build_basic_info(doc: Document, payload: dict[str, Any]) -> None:
    _section_heading(doc, "1. Basic Information")
    _label_value_table(doc, [
        ("Course Title (according to the bylaw):", payload.get("course_title") or payload.get("course") or ""),
        ("Course Code (according to the bylaw):", payload.get("course_code") or ""),
        ("Department/s that participated in the teaching:", payload.get("department") or ""),
        ("Total number of credit hours/points of the course:", payload.get("credit_hours") or ""),
        ("Course Type:", payload.get("course_type") or ""),
        ("The level to which the course was introduced:", payload.get("level") or ""),
        ("Academic Program:", payload.get("program") or ""),
        ("Faculty/Institute:", payload.get("faculty") or ""),
        ("University/Academy:", payload.get("university") or ""),
        ("Name of Course Coordinator:", payload.get("coordinator") or ""),
        ("Date of approval of the course report:", payload.get("approval_date") or ""),
    ])


def _build_data_and_statistics(doc: Document, payload: dict[str, Any], results: dict[str, Any]) -> None:
    _section_heading(doc, "2. Data and Statistics")

    doc.add_paragraph().add_run("Course Instructors").bold = True
    instr_table = doc.add_table(rows=2, cols=4)
    instr_table.style = "Table Grid"
    hdr = instr_table.rows[0].cells
    for i, h in enumerate(["Instructor Name", "Department", "Academic Degree", "Specialty"]):
        _bold(hdr[i], h)
        _shade_cell(hdr[i], _LIGHT_GREY)
    body = instr_table.rows[1].cells
    body[0].text = payload.get("coordinator") or "Not available"
    body[1].text = payload.get("department") or "Not available"
    body[2].text = payload.get("coordinator_degree") or "Not available"
    body[3].text = payload.get("coordinator_specialty") or "Not available"
    _note(doc, "Additional instructors/teaching assistants should be added manually; only the course coordinator is captured by the analytics pipeline.")

    doc.add_paragraph()
    doc.add_paragraph().add_run("Student Assessment Results").bold = True
    _label_value_table(doc, [
        ("Number of students (who started the course):", _fmt_num(results["enrollment"])),
        ("Number of students (who completed the course / sat for the exam):", _fmt_num(results["enrollment"])),
        ("Total number of students who passed the exams successfully:", _fmt_num(results["passed"])),
        ("Percentage of success (out of total students who sat for the final exam):", _fmt_pct(results["success_pct"])),
        ("Total number of students who failed the exams:", _fmt_num(results["fail_count"])),
        ("Percentage of failure (out of total students who took the final exam):", _fmt_pct(results["failure_pct"])),
        ("Estimated GPA:", _fmt_gpa(payload.get("gpa"))),
    ])

    doc.add_paragraph()
    doc.add_paragraph().add_run("Grade Distribution").bold = True
    grade_table = doc.add_table(rows=3, cols=len(GRADE_ORDER) + 1)
    grade_table.style = "Table Grid"
    hdr = grade_table.rows[0].cells
    _bold(hdr[0], "Grade")
    _shade_cell(hdr[0], _LIGHT_GREY)
    for i, g in enumerate(GRADE_ORDER, start=1):
        _bold(hdr[i], g)
        _shade_cell(hdr[i], _LIGHT_GREY)
    counts_row = grade_table.rows[1].cells
    pct_row = grade_table.rows[2].cells
    _bold(counts_row[0], "Number of students")
    _bold(pct_row[0], "Percentage")
    total = results["enrollment"] or 0
    for i, g in enumerate(GRADE_ORDER, start=1):
        cnt = results["grade_counts"].get(g, 0)
        counts_row[i].text = str(cnt)
        pct_row[i].text = f"{(cnt / total * 100):.1f}%" if total > 0 else "0.0%"

    doc.add_paragraph()
    comment = doc.add_paragraph()
    comment.add_run("Commenting on student results and analyzing performance: ").bold = True
    narrative = (
        f"The course had {_fmt_num(results['enrollment'])} graded students, achieving a success rate of "
        f"{_fmt_pct(results['success_pct'])} and an estimated GPA of {_fmt_gpa(payload.get('gpa'))}. "
        f"The failure rate was {_fmt_pct(results['failure_pct'])}"
    )
    risk_score = _safe_float(payload.get("risk_score"))
    if risk_score is not None:
        narrative += f", with a predicted course risk score of {risk_score:.1f}%."
    else:
        narrative += "."
    comment.add_run(narrative)


def _build_quality_evaluation(doc: Document, payload: dict[str, Any]) -> None:
    _section_heading(doc, "3. Course Quality Evaluation - Students' Evaluation of the Course")
    p = doc.add_paragraph()
    p.add_run(
        "Survey-based course quality evaluation (student ratings of scientific content, teaching methods, "
        "facilities, and examinations) is not part of the uploaded grades workbook. Attach the questionnaire "
        "analysis report separately, or use the Survey Dashboard module to generate and merge it into this report."
    )


def _build_student_feedback(doc: Document, payload: dict[str, Any]) -> None:
    _section_heading(doc, "4. Student Feedback")
    _note(doc, "Feedback must include evaluation of: scientific content – teaching and learning methods – facilities and learning resources – examinations.")
    _label_value_table(doc, [
        ("Means of Evaluation:", payload.get("feedback_means") or "Not available"),
        ("Timing of Evaluation:", payload.get("feedback_timing") or "Not available"),
        ("Number of students who participated in the course evaluation:", payload.get("feedback_count") or "Not available"),
        ("Percentage of participants to the total number:", payload.get("feedback_pct") or "Not available"),
        ("Important points of satisfaction:", payload.get("feedback_satisfaction") or "Not available"),
        ("Important points of dissatisfaction:", payload.get("feedback_dissatisfaction") or "Not available"),
    ])


def _build_instructors_reflection(doc: Document, payload: dict[str, Any]) -> None:
    _section_heading(doc, "5. Instructors' Reflection")
    p = doc.add_paragraph()
    p.add_run(payload.get("instructors_reflection") or "Not available — to be completed by the course coordinator.")


def _build_enhancement(doc: Document, payload: dict[str, Any], recommendations: list[str]) -> None:
    _section_heading(doc, "6. Course Enhancement")
    doc.add_paragraph().add_run(
        "Course development plan for the next academic semester/year "
        "(derived from analytics-based recommendations where available):"
    ).bold = True

    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(["No.", "Points that need development", "Corrective/Improvement Actions", "Notes"]):
        _bold(hdr[i], h)
        _shade_cell(hdr[i], _LIGHT_GREY)

    rows_source = recommendations[:6] if recommendations else [
        "Maintain current teaching practices; no high-risk indicators detected in the uploaded data.",
    ]
    for i, rec in enumerate(rows_source, start=1):
        row = table.add_row().cells
        row[0].text = str(i)
        row[1].text = "Course performance"
        row[2].text = rec
        row[3].text = ""


def _build_signatures(doc: Document, payload: dict[str, Any]) -> None:
    doc.add_paragraph()
    table = doc.add_table(rows=2, cols=2)
    table.style = "Table Grid"
    _bold(table.rows[0].cells[0], "Name and Signature\nCourse Coordinator")
    _bold(table.rows[0].cells[1], "Name and Signature\nHead of the Department Council")
    table.rows[1].cells[0].text = payload.get("coordinator") or "Not available"
    table.rows[1].cells[1].text = payload.get("department_head") or "Not available"


# ---------------------------------------------------------------------------
# Public entry point
# ---------------------------------------------------------------------------

def build_course_report_docx(payload: dict[str, Any]) -> bytes:
    """
    payload keys (all optional except 'course'):
      course, course_code, academic_year, semester, department, credit_hours,
      course_type, level, program, faculty, university, coordinator,
      coordinator_degree, coordinator_specialty, approval_date, department_head,
      enrollment, gpa, failure_rate, excellence_rate, risk_score, grade_distribution,
      recommendations (list[str]), feedback_* fields, instructors_reflection.
    """
    import io

    doc = Document()
    style = doc.styles["Normal"]
    style.font.size = Pt(10.5)

    results = _derive_assessment_results(payload)

    _build_title(doc, payload)
    _build_basic_info(doc, payload)
    _build_data_and_statistics(doc, payload, results)
    _build_quality_evaluation(doc, payload)
    _build_student_feedback(doc, payload)
    _build_instructors_reflection(doc, payload)
    _build_enhancement(doc, payload, payload.get("recommendations") or [])
    _build_signatures(doc, payload)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
