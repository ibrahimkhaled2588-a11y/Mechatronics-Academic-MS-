"""Smoke test for backend/ssr_report.py (Phase 7 — final SSR integration).

Seeds a little data across every accreditation-support module, builds the
SSR docx, and checks each standard's section actually surfaces that data
(not just that the file is non-empty). Plain-script convention (see
test_indicators.py) — run directly: `python tests/test_ssr_report.py`.
"""
import io
import os
import sys
import tempfile

backend_dir = os.path.abspath(os.path.join(os.path.dirname(__file__), '..'))
sys.path.append(backend_dir)

_tmp_db = os.path.join(tempfile.gettempdir(), "test_ssr_report.db")
if os.path.exists(_tmp_db):
    os.remove(_tmp_db)
os.environ["ACCREDITATION_DB_PATH"] = _tmp_db

import alumni  # noqa: E402
import curriculum_mapping as cm  # noqa: E402
import faculty_data as fd  # noqa: E402
import governance as gov  # noqa: E402
import indicators  # noqa: E402
import resources as res  # noqa: E402
import ssr_report  # noqa: E402
from docx import Document  # noqa: E402


def _all_docx_text(docx_bytes: bytes) -> str:
    doc = Document(io.BytesIO(docx_bytes))
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.append(cell.text)
    return "\n".join(parts)


# --- seed one thing per module ---
gov.create_mission_version("Prepare graduates for mechatronics engineering practice.")
gov.add_stakeholder_entry("Ahmed Ibrahim", consulted_on="2026-02-01", topic="Curriculum priorities", stakeholder_role="Alumni")

ilo1 = cm.create_ilo("Apply engineering fundamentals", "ILO1")
cm.create_course("MEC 305 / Control Systems")

alumni.create_alumnus("Sara Mostafa", graduation_year=2023, employer="Delta Electronics")

f1 = fd.create_faculty("Dr. A", specialization="Control Systems")
fd.create_teaching_load(f1["id"], "Fall 2025", "Control Systems", 6)

res.create_equipment("CNC Machine", location="Lab 3", next_maintenance_date="2020-01-01")  # overdue
res.create_budget_entry("2025-2026", "Lab Equipment", 50000)

indicators.seed_defaults()
std7_first = indicators.list_indicators(standard_number=7)[0]
indicators.add_log_entry(std7_first["id"], weakness_identified="No self-study report existed", action_taken="Built the SSR generator", entry_status="complete")

# --- build without analysis (Standard 3 should note missing data) ---
docx_bytes = ssr_report.build_ssr_docx(analysis=None)
assert isinstance(docx_bytes, bytes) and len(docx_bytes) > 0
text = _all_docx_text(docx_bytes)

assert "Self-Study Report" in text
assert "Prepare graduates for mechatronics engineering practice." in text, "Standard 1 mission text missing"
assert "Ahmed Ibrahim" in text, "Standard 1 stakeholder log missing"
assert "Apply engineering fundamentals" in text, "Standard 2 ILO missing"
assert "No analytics data was supplied" in text, "Standard 3 should note missing analysis"
assert "Sara Mostafa" not in text, "alumni names aren't listed individually, only aggregate stats"
assert "Delta Electronics" not in text  # aggregate only, not a name/employer table
assert "1" in text and "Faculty members" in text, "Standard 5 KPI table missing"
assert "CNC Machine" in text, "Standard 6 maintenance-due table missing"
assert "No self-study report existed" in text, "Standard 7 closing-the-loop entry missing"
assert "Built the SSR generator" in text

# --- build with a minimal analysis payload (Standard 3 KPIs should appear) ---
analysis = {
    "academic_analytics": {
        "Sheet1": {
            "all_courses": [
                {"course": "MEC 305 / Control Systems", "enrollment": 30, "failure_rate": 15.0, "gpa_estimate": 3.1, "excellence_rate": 20.0},
            ]
        }
    },
    "metadata": {"kpis": {}},
}
docx_bytes2 = ssr_report.build_ssr_docx(analysis=analysis)
text2 = _all_docx_text(docx_bytes2)
assert "No analytics data was supplied" not in text2
assert "Average GPA" in text2
assert "Courses analyzed" in text2

os.remove(_tmp_db)
print("All ssr_report.py tests passed.")
