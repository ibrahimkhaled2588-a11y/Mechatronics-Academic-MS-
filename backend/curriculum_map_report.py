"""
Curriculum Map DOCX builder (Standard 2 — Program Design).

Follows the same from-scratch python-docx pattern as course_report_docx.py:
a courses x ILOs coverage matrix table, plus a findings section (zero/low
coverage ILOs, heavy-duplication ILOs, courses with no mapped ILOs) derived
from curriculum_mapping.compute_coverage_summary().
"""
from __future__ import annotations

from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

_HEADER_BLUE = RGBColor(0x1E, 0x3A, 0x5F)
_LIGHT_GREY = "D9D9D9"
_CHECK_GREEN = "DCFCE7"
_WARN_RED = "FEE2E2"


def _shade_cell(cell, hex_color: str) -> None:
    shd = cell._tc.get_or_add_tcPr().makeelement(qn("w:shd"), {
        qn("w:val"): "clear", qn("w:color"): "auto", qn("w:fill"): hex_color,
    })
    cell._tc.get_or_add_tcPr().append(shd)


def _bold(cell, text: str, size: int | None = None) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.bold = True
    if size is not None:
        run.font.size = Pt(size)


def _section_heading(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(13)
    run.font.color.rgb = _HEADER_BLUE


def _bullet_list(doc: Document, items: list[str], empty_text: str) -> None:
    if not items:
        doc.add_paragraph().add_run(empty_text).italic = True
        return
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def build_curriculum_map_docx(data: dict[str, Any]) -> bytes:
    """
    data: output of curriculum_mapping.get_export_data(), i.e.
      { ilos: [...], courses: [...], matrix: {course_id: {ilo_id: bool}}, summary: {...} }
    """
    import io

    ilos: list[dict] = data.get("ilos") or []
    courses: list[dict] = data.get("courses") or []
    matrix: dict = data.get("matrix") or {}
    summary: dict = data.get("summary") or {}

    doc = Document()
    style = doc.styles["Normal"]
    style.font.size = Pt(10)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Curriculum Map — Courses x Program ILOs")
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = _HEADER_BLUE

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.add_run(f"{len(courses)} courses  |  {len(ilos)} Program Intended Learning Outcomes").italic = True

    _section_heading(doc, "Coverage Matrix")
    if not ilos or not courses:
        doc.add_paragraph().add_run(
            "No ILOs and/or courses have been entered yet — add them in the Curriculum Mapping page before exporting."
        ).italic = True
    else:
        table = doc.add_table(rows=1, cols=len(ilos) + 1)
        table.style = "Table Grid"
        hdr = table.rows[0].cells
        _bold(hdr[0], "Course")
        _shade_cell(hdr[0], _LIGHT_GREY)
        for i, ilo in enumerate(ilos, start=1):
            label = ilo.get("ilo_code") or f"ILO{ilo['id']}"
            _bold(hdr[i], label, size=9)
            _shade_cell(hdr[i], _LIGHT_GREY)

        for course in courses:
            row = table.add_row().cells
            row[0].text = course.get("course_name", "")
            course_map = matrix.get(course["id"], {})
            for i, ilo in enumerate(ilos, start=1):
                mapped = course_map.get(ilo["id"], False)
                row[i].text = "✓" if mapped else ""
                _shade_cell(row[i], _CHECK_GREEN if mapped else "FFFFFF")

    _section_heading(doc, "Findings")

    doc.add_paragraph().add_run("ILOs with zero coverage (no course addresses them):").bold = True
    _bullet_list(
        doc,
        [f"{i.get('ilo_code') or ('ILO' + str(i['id']))}: {i['ilo_text']}" for i in summary.get("zero_coverage_ilos", [])],
        "None — every ILO is addressed by at least one course.",
    )

    doc.add_paragraph().add_run("ILOs with low coverage:").bold = True
    _bullet_list(
        doc,
        [f"{i.get('ilo_code') or ('ILO' + str(i['id']))}: {i['ilo_text']}" for i in summary.get("low_coverage_ilos", [])],
        "None.",
    )

    doc.add_paragraph().add_run("ILOs with heavy duplication (mapped by many courses — possible redundancy):").bold = True
    _bullet_list(
        doc,
        [f"{i.get('ilo_code') or ('ILO' + str(i['id']))}: {i['ilo_text']}" for i in summary.get("heavy_duplication_ilos", [])],
        "None.",
    )

    doc.add_paragraph().add_run("Courses with no mapped ILOs:").bold = True
    _bullet_list(
        doc,
        [c["course_name"] for c in summary.get("courses_without_ilos", [])],
        "None — every course is mapped to at least one ILO.",
    )

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
